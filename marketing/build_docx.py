"""
Build vpma_marketing_strategy.docx from the markdown source.
Run from repo root: backend/.venv/bin/python marketing/build_docx.py
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Palette ────────────────────────────────────────────────────────────────
VERA_PURPLE   = RGBColor(0x6C, 0x63, 0xFF)
VERA_DARK     = RGBColor(0x1A, 0x17, 0x2E)
VERA_ACCENT   = RGBColor(0xF0, 0xA5, 0x00)
TEXT_BODY     = RGBColor(0x1A, 0x17, 0x2E)
TEXT_MUTED    = RGBColor(0x5A, 0x56, 0x7A)
RULE_GRAY     = RGBColor(0xE0, 0xDE, 0xF4)
TABLE_HEADER  = RGBColor(0x6C, 0x63, 0xFF)
TABLE_ALT     = RGBColor(0xF4, 0xF3, 0xFF)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)

# ─── Document setup ──────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)

# ─── Style helpers ───────────────────────────────────────────────────────────
def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def set_cell_borders(table, color: RGBColor = RULE_GRAY):
    """Apply thin borders to all cells."""
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), hex_color)
        tblBorders.append(el)
    tblPr.append(tblBorders)

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), f"{RULE_GRAY[0]:02X}{RULE_GRAY[1]:02X}{RULE_GRAY[2]:02X}")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

def para_fmt(p, space_before=0, space_after=6, line_spacing=None):
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if line_spacing:
        p.paragraph_format.line_spacing = Pt(line_spacing)

def apply_inline(run, bold=False, italic=False, color=None, size=None, font_name=None):
    run.bold   = bold
    run.italic = italic
    if color:   run.font.color.rgb = color
    if size:    run.font.size = Pt(size)
    if font_name: run.font.name = font_name

def add_cover_page(doc):
    """Full-color cover page."""
    # Cover heading
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_fmt(p_title, space_before=60, space_after=4)
    run = p_title.add_run("VetAgent (VPMA)")
    apply_inline(run, bold=True, color=VERA_PURPLE, size=32, font_name="Calibri")

    p_sub = doc.add_paragraph()
    para_fmt(p_sub, space_before=0, space_after=6)
    run = p_sub.add_run("Marketing Strategy")
    apply_inline(run, bold=True, color=VERA_DARK, size=22, font_name="Calibri")

    p_tag = doc.add_paragraph()
    para_fmt(p_tag, space_before=10, space_after=20)
    run = p_tag.add_run("Meet Vera. Your practice's AI Chief of Staff.")
    apply_inline(run, italic=True, color=TEXT_MUTED, size=13, font_name="Calibri")

    add_horizontal_rule(doc)

    meta = doc.add_paragraph()
    para_fmt(meta, space_before=12, space_after=4)
    run = meta.add_run("Version 1.0  ·  June 2026  ·  Status: Approved for Execution")
    apply_inline(run, color=TEXT_MUTED, size=10, font_name="Calibri")

    doc.add_page_break()

def add_h1(doc, text):
    p = doc.add_paragraph()
    para_fmt(p, space_before=20, space_after=6)
    run = p.add_run(text.upper())
    apply_inline(run, bold=True, color=VERA_PURPLE, size=13, font_name="Calibri")
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), f"{VERA_PURPLE[0]:02X}{VERA_PURPLE[1]:02X}{VERA_PURPLE[2]:02X}")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    para_fmt(p, space_before=14, space_after=4)
    run = p.add_run(text)
    apply_inline(run, bold=True, color=VERA_DARK, size=12, font_name="Calibri")
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    para_fmt(p, space_before=10, space_after=3)
    run = p.add_run(text)
    apply_inline(run, bold=True, color=VERA_ACCENT, size=11, font_name="Calibri")
    return p

def add_body(doc, text, italic=False):
    """Add a body paragraph, handling **bold** and *italic* inline markdown."""
    if not text.strip():
        return
    p = doc.add_paragraph()
    para_fmt(p, space_before=0, space_after=5, line_spacing=14)
    _add_inline_text(p, text, base_italic=italic, base_color=TEXT_BODY)
    return p

def add_callout(doc, text):
    """Indented callout block (blockquote)."""
    p = doc.add_paragraph()
    para_fmt(p, space_before=4, space_after=4, line_spacing=14)
    p.paragraph_format.left_indent = Inches(0.35)
    # Left bar
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), f"{VERA_PURPLE[0]:02X}{VERA_PURPLE[1]:02X}{VERA_PURPLE[2]:02X}")
    pBdr.append(left)
    pPr.append(pBdr)
    _add_inline_text(p, text, base_italic=True, base_color=TEXT_MUTED, base_size=10.5)
    return p

def add_code_block(doc, lines):
    """Monospaced code/narrative block."""
    for line in lines:
        p = doc.add_paragraph()
        para_fmt(p, space_before=0, space_after=2)
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(line)
        apply_inline(run, color=VERA_DARK, size=9.5, font_name="Courier New")

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph()
    para_fmt(p, space_before=1, space_after=3, line_spacing=13)
    p.paragraph_format.left_indent  = Inches(0.3 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    # Bullet character
    bullet_run = p.add_run("•  ")
    apply_inline(bullet_run, bold=True, color=VERA_PURPLE, size=10)
    _add_inline_text(p, text, base_color=TEXT_BODY)
    return p

def add_checklist_item(doc, text):
    p = doc.add_paragraph()
    para_fmt(p, space_before=2, space_after=3)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    box_run = p.add_run("☐  ")
    apply_inline(box_run, color=VERA_PURPLE, size=10)
    _add_inline_text(p, text, base_color=TEXT_BODY)

def _add_inline_text(p, text, base_italic=False, base_color=TEXT_BODY, base_size=10.5):
    """Parse **bold** and *italic* markdown inline in a single paragraph."""
    # Split on bold (**...**) and italic (*...*)
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            apply_inline(run, bold=True, color=base_color, size=base_size)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = p.add_run(part[1:-1])
            apply_inline(run, italic=True, color=base_color, size=base_size)
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            apply_inline(run, color=VERA_PURPLE, size=9.5, font_name="Courier New")
        else:
            run = p.add_run(part)
            apply_inline(run, italic=base_italic, color=base_color, size=base_size)

def add_markdown_table(doc, header_row, data_rows):
    """Render a markdown table as a styled DOCX table."""
    col_count = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, cell_text in enumerate(header_row):
        hdr_cells[i].text = cell_text.strip()
        set_cell_bg(hdr_cells[i], TABLE_HEADER)
        run = hdr_cells[i].paragraphs[0].runs
        if run:
            run[0].bold = True
            run[0].font.color.rgb = WHITE
            run[0].font.size = Pt(9.5)
            run[0].font.name = "Calibri"
        hdr_cells[i].paragraphs[0].paragraph_format.space_before = Pt(3)
        hdr_cells[i].paragraphs[0].paragraph_format.space_after = Pt(3)

    # Data rows
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.rows[row_idx + 1].cells
        bg = TABLE_ALT if row_idx % 2 == 0 else WHITE
        for j, cell_text in enumerate(row_data):
            clean = re.sub(r'\*\*|~~|__', '', cell_text.strip())
            row_cells[j].text = clean
            set_cell_bg(row_cells[j], bg)
            if row_cells[j].paragraphs[0].runs:
                row_cells[j].paragraphs[0].runs[0].font.size = Pt(9.5)
                row_cells[j].paragraphs[0].runs[0].font.name = "Calibri"
                row_cells[j].paragraphs[0].runs[0].font.color.rgb = TEXT_BODY
            row_cells[j].paragraphs[0].paragraph_format.space_before = Pt(2)
            row_cells[j].paragraphs[0].paragraph_format.space_after = Pt(2)

    set_cell_borders(table)
    doc.add_paragraph()  # breathing room after table

# ─── Markdown → DOCX parser ──────────────────────────────────────────────────
def parse_table_line(line):
    """Split a markdown table row into cells."""
    cells = [c.strip() for c in line.strip().strip('|').split('|')]
    return cells

def is_table_separator(line):
    return bool(re.match(r'^\|?\s*[-:]+\s*(\|\s*[-:]+\s*)*\|?\s*$', line.strip()))

def is_table_row(line):
    return '|' in line and not line.strip().startswith('#')

def convert():
    md_path = Path(__file__).parent / "vpma_marketing_strategy.md"
    out_path = Path(__file__).parent / "VetAgent_Marketing_Strategy.docx"

    lines = md_path.read_text(encoding='utf-8').splitlines()

    add_cover_page(doc)

    i = 0
    in_code_block = False
    code_lines = []
    table_buffer = []  # [(header_row, [data_rows])]
    collecting_table = False
    table_header = None
    table_data = []

    def flush_table():
        nonlocal table_header, table_data, collecting_table
        if table_header and table_data:
            add_markdown_table(doc, table_header, table_data)
        table_header = None
        table_data = []
        collecting_table = False

    while i < len(lines):
        line = lines[i]
        raw = line.rstrip()

        # ── Code block ──────────────────────────────────────────────────────
        if raw.startswith('```'):
            if in_code_block:
                # End of code block
                if code_lines:
                    add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                flush_table()
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(raw)
            i += 1
            continue

        # ── Table handling ───────────────────────────────────────────────────
        if is_table_row(raw):
            if is_table_separator(raw):
                # separator line — skip
                i += 1
                continue
            cells = parse_table_line(raw)
            if not collecting_table:
                # First row = header
                flush_table()
                table_header = cells
                collecting_table = True
                table_data = []
            else:
                table_data.append(cells)
            i += 1
            continue
        else:
            if collecting_table:
                flush_table()

        # ── Horizontal rule ──────────────────────────────────────────────────
        if raw.strip() in ('---', '***', '___'):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Headings ─────────────────────────────────────────────────────────
        if raw.startswith('# ') and not raw.startswith('## '):
            text = raw[2:].strip()
            # Skip title — already on cover page
            if 'VetAgent (VPMA) Marketing Strategy' not in text:
                add_h1(doc, text)
            i += 1
            continue

        if raw.startswith('## '):
            text = raw[3:].strip()
            # Strip leading "N. " numbering
            text = re.sub(r'^\d+\.\s+', '', text)
            add_h1(doc, text)
            i += 1
            continue

        if raw.startswith('### '):
            text = raw[4:].strip()
            add_h2(doc, text)
            i += 1
            continue

        if raw.startswith('#### '):
            text = raw[5:].strip()
            add_h3(doc, text)
            i += 1
            continue

        # ── Blockquote ───────────────────────────────────────────────────────
        if raw.startswith('> '):
            text = raw[2:].strip()
            add_callout(doc, text)
            i += 1
            continue

        # ── Bullets ──────────────────────────────────────────────────────────
        bullet_match = re.match(r'^(\s*)[-*]\s+(.*)', raw)
        if bullet_match:
            indent = len(bullet_match.group(1)) // 2
            text = bullet_match.group(2).strip()
            add_bullet(doc, text, level=indent)
            i += 1
            continue

        # ── Numbered list ────────────────────────────────────────────────────
        num_match = re.match(r'^(\s*)\d+\.\s+(.*)', raw)
        if num_match:
            indent = len(num_match.group(1)) // 2
            text = num_match.group(2).strip()
            add_bullet(doc, text, level=indent)
            i += 1
            continue

        # ── Checklist ────────────────────────────────────────────────────────
        check_match = re.match(r'^- \[ \]\s+(.*)', raw)
        if check_match:
            add_checklist_item(doc, check_match.group(1))
            i += 1
            continue

        # ── Bold-only line (used for sub-section labels) ──────────────────
        bold_only = re.match(r'^\*\*(.+)\*\*$', raw.strip())
        if bold_only:
            p = doc.add_paragraph()
            para_fmt(p, space_before=8, space_after=2)
            run = p.add_run(bold_only.group(1))
            apply_inline(run, bold=True, color=VERA_DARK, size=10.5, font_name="Calibri")
            i += 1
            continue

        # ── Meta line (Version / Prepared for) ──────────────────────────────
        if raw.startswith('**Version:**') or raw.startswith('**Prepared for:**'):
            p = doc.add_paragraph()
            para_fmt(p, space_before=0, space_after=2)
            _add_inline_text(p, raw, base_color=TEXT_MUTED, base_size=9.5)
            i += 1
            continue

        # ── Empty line ───────────────────────────────────────────────────────
        if not raw.strip():
            i += 1
            continue

        # ── Body paragraph ───────────────────────────────────────────────────
        add_body(doc, raw)
        i += 1

    flush_table()

    # Save
    doc.save(str(out_path))
    print(f"✅ Saved: {out_path}")
    return out_path

if __name__ == "__main__":
    convert()
